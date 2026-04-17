import os
import io
import json
import re
import subprocess
import tempfile
import shutil
import base64
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware

import docx
from docx import Document as DocxDocument
from docxtpl import DocxTemplate
from PIL import Image
import httpx

app = FastAPI(title="DocFlow API", version="3.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

SUPABASE_URL = os.getenv("SUPABASE_URL", "")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY", "")
STORAGE_BUCKET = os.getenv("STORAGE_BUCKET", "generated")
TEMPLATES_BUCKET = os.getenv("TEMPLATES_BUCKET", "templates")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")


# ─── HEALTH ───────────────────────────────────────────────

@app.get("/health")
def health():
    return {"status": "ok"}


# ─── EXTRACT TEXT ─────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_text_from_pdf_fast(file_bytes: bytes):
    """Fast path: PyMuPDF text extraction only. Returns (text, page_count)."""
    import fitz
    parts = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page_count = len(doc)
    for page in doc:
        text = page.get_text()
        if text and text.strip():
            parts.append(text.strip())
    doc.close()
    return "\n".join(parts), page_count


def alphanumeric_ratio(text: str) -> float:
    """Fraction of non-whitespace characters that are alphanumeric.
    Garbled OCR output has a low ratio due to excess punctuation/symbols."""
    non_ws = [c for c in text if not c.isspace()]
    if not non_ws:
        return 0.0
    alnum = sum(1 for c in non_ws if c.isalnum())
    return alnum / len(non_ws)


async def ocr_pdf_with_gemini(file_bytes: bytes) -> str:
    """OCR via Gemini Vision API. Sends the PDF directly — no image conversion."""

    b64_pdf = base64.b64encode(file_bytes).decode("utf-8")

    parts = [
        {
            "inline_data": {
                "mime_type": "application/pdf",
                "data": b64_pdf
            }
        },
        {
            "text": (
                "Extraia TODO o texto visivel neste documento PDF. "
                "Retorne apenas o texto extraido, na ordem em que aparece, "
                "sem comentarios, sem formatacao markdown, sem explicacoes."
            )
        }
    ]

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}"

    async with httpx.AsyncClient(timeout=180) as client:
        resp = await client.post(url, json={
            "contents": [{"parts": parts}],
            "generationConfig": {
                "maxOutputTokens": 65536
            }
        })

    if resp.status_code != 200:
        raise Exception(f"Gemini API error: {resp.status_code}")

    data = resp.json()
    candidates = data.get("candidates", [])
    if not candidates:
        raise Exception("Gemini retornou sem candidates")

    text_parts = []
    for part in candidates[0].get("content", {}).get("parts", []):
        if "text" in part:
            text_parts.append(part["text"])

    return "\n".join(text_parts).strip()


def ocr_pdf_with_tesseract(file_bytes: bytes) -> str:
    """Fallback OCR: Tesseract, todas as paginas, 72 DPI, paralelo."""
    import fitz
    import pytesseract

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    images = []
    for page in doc:
        pix = page.get_pixmap(dpi=72)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    doc.close()

    def ocr_image(img):
        text = pytesseract.image_to_string(img, lang="por")
        clean = text.strip()
        if not clean or len(clean) < 20:
            return None
        return clean

    with ThreadPoolExecutor(max_workers=4) as executor:
        results = list(executor.map(ocr_image, images))

    ocr_parts = [r for r in results if r]
    return "\n".join(ocr_parts)


def extract_text_from_image(file_bytes: bytes) -> str:
    try:
        import pytesseract
        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image, lang="por")
        return text.strip()
    except Exception as e:
        return f"[OCR falhou: {str(e)}]"


@app.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    content = await file.read()
    filename = (file.filename or "").lower()

    page_count = None
    ocr_triggered = False
    ocr_reason = None
    pre_ocr_chars = None
    pre_ocr_alnum_ratio = None

    if filename.endswith(".docx"):
        text = extract_text_from_docx(content)
    elif filename.endswith(".pdf"):
        text, page_count = extract_text_from_pdf_fast(content)
        pre_ocr_chars = len(text)
        pre_ocr_alnum_ratio = alphanumeric_ratio(text)
        chars_per_page = (pre_ocr_chars / page_count) if page_count else 0

        if chars_per_page < 150:
            ocr_triggered = True
            ocr_reason = f"chars_per_page={chars_per_page:.0f} < 150"
        elif pre_ocr_alnum_ratio < 0.60:
            ocr_triggered = True
            ocr_reason = f"alphanumeric_ratio={pre_ocr_alnum_ratio:.2f} < 0.60"

        if ocr_triggered:
            if GEMINI_API_KEY:
                try:
                    text = await ocr_pdf_with_gemini(content)
                except Exception:
                    text = ocr_pdf_with_tesseract(content)
            else:
                text = ocr_pdf_with_tesseract(content)
    elif filename.endswith((".jpg", ".jpeg", ".png", ".webp", ".tiff", ".bmp")):
        text = extract_text_from_image(content)
    else:
        try:
            text, page_count = extract_text_from_pdf_fast(content)
        except Exception:
            try:
                text = extract_text_from_docx(content)
            except Exception:
                raise HTTPException(400, "Formato nao suportado. Use DOCX, PDF ou imagem.")

    return {
        "text": text,
        "char_count": len(text),
        "page_count": page_count,
        "ocr_triggered": ocr_triggered,
        "ocr_reason": ocr_reason,
        "pre_ocr_chars": pre_ocr_chars,
        "pre_ocr_alnum_ratio": round(pre_ocr_alnum_ratio, 3) if pre_ocr_alnum_ratio is not None else None,
    }


# ─── PDF TO DOCX CONVERSION ──────────────────────────────

def pdf_to_docx(pdf_path: str, output_dir: str) -> str:
    result = subprocess.run(
        ["libreoffice", "--headless", "--infilter=writer_pdf_import", "--convert-to", "docx", "--outdir", output_dir, pdf_path],
        capture_output=True, text=True, timeout=120,
    )
    docx_name = Path(pdf_path).stem + ".docx"
    docx_path = os.path.join(output_dir, docx_name)
    if not os.path.exists(docx_path):
        raise HTTPException(500, f"Conversao PDF para DOCX falhou: {result.stderr}")
    return docx_path


def ensure_docx(file_path: str, tmpdir: str) -> str:
    if file_path.lower().endswith(".pdf"):
        return pdf_to_docx(file_path, tmpdir)
    return file_path


# ─── PLACEHOLDER FIXES ────────────────────────────────────

def merge_fragmented_placeholders(docx_path: str):
    doc = DocxDocument(docx_path)

    def fix_paragraph(para):
        full_text = "".join(run.text for run in para.runs)
        if "{{" not in full_text:
            return
        if para.runs:
            para.runs[0].text = full_text
            for run in para.runs[1:]:
                run.text = ""

    for para in doc.paragraphs:
        fix_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fix_paragraph(para)

    for section in doc.sections:
        for part in [section.header, section.footer,
                     section.first_page_header, section.first_page_footer]:
            if part:
                for para in part.paragraphs:
                    fix_paragraph(para)

    doc.save(docx_path)


def normalize_placeholders(docx_path: str):
    doc = DocxDocument(docx_path)
    pattern = re.compile(r"\{\{(\w+)\}\}")

    def fix_runs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                if "{{" in run.text:
                    run.text = pattern.sub(r"{{ \1 }}", run.text)

    fix_runs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fix_runs(cell.paragraphs)

    for section in doc.sections:
        for part in [section.header, section.footer,
                     section.first_page_header, section.first_page_footer]:
            if part:
                fix_runs(part.paragraphs)

    doc.save(docx_path)


# ─── SUPABASE STORAGE ─────────────────────────────────────

async def ensure_bucket(bucket: str):
    try:
        create_url = f"{SUPABASE_URL}/storage/v1/bucket"
        headers = {
            "Authorization": f"Bearer {SUPABASE_KEY}",
            "Content-Type": "application/json",
        }
        async with httpx.AsyncClient(timeout=10) as client:
            await client.post(create_url, json={"id": bucket, "name": bucket, "public": True}, headers=headers)
    except Exception:
        pass


async def upload_to_supabase(file_bytes: bytes, path: str, content_type: str, bucket: str = None) -> str:
    bucket = bucket or STORAGE_BUCKET
    if not SUPABASE_URL or not SUPABASE_KEY:
        raise HTTPException(500, "Supabase nao configurado")

    url = f"{SUPABASE_URL}/storage/v1/object/{bucket}/{path}"
    headers = {
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": content_type,
        "x-upsert": "true",
    }
    async with httpx.AsyncClient(timeout=60) as client:
        resp = await client.post(url, content=file_bytes, headers=headers)
        if resp.status_code not in (200, 201):
            raise HTTPException(500, f"Upload Supabase falhou: {resp.text}")

    public_url = f"{SUPABASE_URL}/storage/v1/object/public/{bucket}/{path}"
    return public_url


# ─── APPLY MARKERS ────────────────────────────────────────

def apply_markers_to_docx(docx_path: str, fields: list) -> tuple:
    doc = DocxDocument(docx_path)

    replacements = []
    for field in fields:
        if not field.get("confirmed", False):
            continue
        if field.get("is_fixed", False):
            continue
        original = field.get("original_value", "")
        code = field.get("code", "")
        if original and original.strip() and code:
            replacements.append((original.strip(), "{{" + code + "}}"))

    replacements.sort(key=lambda x: len(x[0]), reverse=True)

    debug_log = []
    count = 0

    def replace_in_paragraph(paragraph):
        nonlocal count
        for old_text, new_text in replacements:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
                    debug_log.append(f"RUN: [{old_text[:30]}] -> [{new_text}]")

        full_text = "".join(run.text for run in paragraph.runs)
        needs_merge = False
        for old_text, new_text in replacements:
            if old_text in full_text and "{{" + old_text.split("{{")[-1] not in full_text:
                needs_merge = True
                break

        if needs_merge:
            merged = "".join(run.text for run in paragraph.runs)
            for old_text, new_text in replacements:
                if old_text in merged:
                    merged = merged.replace(old_text, new_text)
                    count += 1
                    debug_log.append(f"MERGE: [{old_text[:30]}] -> [{new_text}]")
            if paragraph.runs:
                paragraph.runs[0].text = merged
                for run in paragraph.runs[1:]:
                    run.text = ""

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            replace_in_paragraph(para)

    process_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    for section in doc.sections:
        for part in [section.header, section.footer,
                     section.first_page_header, section.first_page_footer]:
            if part:
                process_paragraphs(part.paragraphs)

    output_path = docx_path.replace(".docx", "_marked.docx")
    doc.save(output_path)
    return output_path, len(replacements), count, len(fields), debug_log


@app.post("/apply-markers")
async def apply_markers(
    template: UploadFile = File(...),
    fields: str = Form(...),
    template_id: str = Form(""),
    project_id: str = Form(""),
):
    try:
        fields_list = json.loads(fields)
    except json.JSONDecodeError:
        raise HTTPException(400, "JSON invalido no campo 'fields'")

    template_bytes = await template.read()
    template_filename = template.filename or "template.docx"

    tmpdir = tempfile.mkdtemp()
    try:
        template_path = os.path.join(tmpdir, template_filename)
        with open(template_path, "wb") as f:
            f.write(template_bytes)

        docx_path = ensure_docx(template_path, tmpdir)
        marked_path, num_rules, num_replacements, num_fields, debug_log = apply_markers_to_docx(docx_path, fields_list)

        with open(marked_path, "rb") as f:
            marked_bytes = f.read()

        result = {
            "template_id": template_id,
            "debug_num_fields": num_fields,
            "debug_num_rules": num_rules,
            "debug_num_replacements": num_replacements,
            "debug_log": debug_log[:30],
        }

        if SUPABASE_URL and SUPABASE_KEY:
            await ensure_bucket(TEMPLATES_BUCKET)
            storage_path = f"{project_id}/{template_id}/marked.docx"
            result["marked_url"] = await upload_to_supabase(
                marked_bytes, storage_path,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                bucket=TEMPLATES_BUCKET,
            )
            with open(docx_path, "rb") as f:
                original_bytes = f.read()
            original_path = f"{project_id}/{template_id}/original.docx"
            result["original_url"] = await upload_to_supabase(
                original_bytes, original_path,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                bucket=TEMPLATES_BUCKET,
            )
        else:
            result["marked_base64"] = base64.b64encode(marked_bytes).decode("utf-8")

        return result

    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


# ─── GENERATE DOCUMENT ────────────────────────────────────

def docx_to_pdf(docx_path: str, output_dir: str) -> str:
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        raise HTTPException(500, f"Erro ao converter para PDF: {result.stderr}")

    pdf_name = Path(docx_path).stem + ".pdf"
    pdf_path = os.path.join(output_dir, pdf_name)
    if not os.path.exists(pdf_path):
        raise HTTPException(500, "PDF nao foi gerado")
    return pdf_path


@app.post("/generate")
async def generate(
    template: UploadFile = File(...),
    data: str = Form(...),
    formats: str = Form("pdf,docx"),
    generation_id: str = Form(""),
    file_prefix: str = Form("documento"),
):
    try:
        replacements = json.loads(data)
    except json.JSONDecodeError:
        raise HTTPException(400, "JSON invalido no campo 'data'")

    requested_formats = [f.strip() for f in formats.split(",")]
    template_bytes = await template.read()
    template_filename = template.filename or "template.docx"

    tmpdir = tempfile.mkdtemp()
    try:
        template_path = os.path.join(tmpdir, template_filename)
        with open(template_path, "wb") as f:
            f.write(template_bytes)

        template_path = ensure_docx(template_path, tmpdir)

        merge_fragmented_placeholders(template_path)
        normalize_placeholders(template_path)

        tpl = DocxTemplate(template_path)
        tpl.render(replacements)
        output_docx = os.path.join(tmpdir, f"output_{Path(template_path).name}")
        tpl.save(output_docx)

        result = {"file_name": file_prefix}

        if "docx" in requested_formats:
            with open(output_docx, "rb") as f:
                docx_bytes = f.read()
            if SUPABASE_URL and SUPABASE_KEY:
                path = f"{generation_id}/{file_prefix}.docx"
                result["docx_url"] = await upload_to_supabase(
                    docx_bytes, path,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                result["docx_base64"] = base64.b64encode(docx_bytes).decode("utf-8")

        if "pdf" in requested_formats:
            pdf_path = docx_to_pdf(output_docx, tmpdir)
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            if SUPABASE_URL and SUPABASE_KEY:
                path = f"{generation_id}/{file_prefix}.pdf"
                result["pdf_url"] = await upload_to_supabase(
                    pdf_bytes, path, "application/pdf"
                )
            else:
                result["pdf_base64"] = base64.b64encode(pdf_bytes).decode("utf-8")

        return result

    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)
